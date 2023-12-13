import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import re
from config import settings
import openai

def fetch_prompt_message(filepath: str) -> str:
    with open(filepath, "r") as f:
        data = f.read()
    return data

def read_docx(docx_path: str) -> str:
    # Load the Document
    doc = docx.Document(docx_path)
    document_text = ""
    # Get all the paragraphs and tables elements
    document_paragraph = doc.paragraphs
    document_table = doc.tables
    body_location = []
    # Get the order of the elements
    for element in doc.element.body:
        if isinstance(element, CT_Tbl):
            body_location.append("T")
        elif isinstance(element, CT_P):
            body_location.append("P")
    # Fetch the document data in the correct order of elements
    para, tab = 0, 0
    for loc in body_location:
        if loc == "P":
            document_text += document_paragraph[para].text + "\n"
            para += 1
        elif loc == "T":
            for row in document_table[tab].rows:
                row_text = ""
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text
                        row_text += paragraph_text + " | "
                document_text += row_text.strip() + "\n"
            tab += 1
    document_text = document_text.strip()
    return document_text

def case_study_extract(evaluation_info: str) -> tuple:
    # Instructions
    important_notice = r"(?i)Important Notice[\s\S]+?Please Note"
    instructions_search = re.search(important_notice, evaluation_info)
    instructions = ""
    if instructions_search is not None:
        instructions = instructions_search.group(0)[17:-11].strip()
    # Abbreviations
    abb_pattern = r"(?i)Abbreviations[\s\S]+?From:"
    abb_search = re.search(abb_pattern, evaluation_info)
    abbreviations = ""
    if abb_search is not None:
        abbreviations = abb_search.group(0)[:-5].strip()
    # Email and Content
    email_pattern = r"Subject:[\s\S]+?([T/t]hanks?|Head of Unit)"
    email_search = re.search(email_pattern, evaluation_info)
    email = ""
    content = evaluation_info
    if email_search is not None:
        email = email_search.group(0).strip()
        content = evaluation_info[email_search.end():].strip()
    
    return instructions, abbreviations, email, content