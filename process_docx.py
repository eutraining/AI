import docx
import os

def create_docx(score: int, summary_text: str, evaluation_text: str) -> None:
    """
    Create a DOCX file with an evaluation report.

    Args:
    score (int): The score to include in the report.
    summary_text (str): The summary text of the evaluation.
    evaluation_text (str): The detailed evaluation text.
    """
    # Create a new document
    doc = docx.Document()

    # Add content to the document
    doc.add_heading('Evaluation report', 0)
    doc.add_paragraph(f'Score: {score}')
    doc.add_paragraph('Summary:')
    doc.add_paragraph(summary_text + "\n\n")
    doc.add_paragraph(evaluation_text)

    # Ensure the evaluations directory exists
    os.makedirs("evaluations", exist_ok=True)

    # Generate a unique file name
    file_name = get_next_available_filename('evaluation_report.docx', "evaluations")

    # Save the document
    doc.save(os.path.join("evaluations", file_name))

def get_next_available_filename(base_filename, directory="."):
    """
    Generate a unique file name by appending a number if a file with the same name already exists.

    Args:
    base_filename (str): The base name of the file.
    directory (str): The directory in which to check for the file. Defaults to the current directory.

    Returns:
    str: A unique file name.
    """
    base, extension = os.path.splitext(base_filename)
    count = 1

    next_filename = f"{base}_{count}{extension}"
    full_path = os.path.join(directory, next_filename)

    while os.path.exists(full_path):
        count += 1
        next_filename = f"{base}_{count}{extension}"
        full_path = os.path.join(directory, next_filename)

    return next_filename