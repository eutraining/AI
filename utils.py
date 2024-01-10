import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import re

def fetch_prompt_message(filepath: str) -> str:
    """Read and return the content of a text file."""
    with open(filepath, "r") as f:
        return f.read()

def load_doc(filepath: str) -> str:
    """Load a DOCX file and return its text content."""
    doc = docx.Document(filepath)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

def read_docx(docx_path: str) -> str:
    """
    Read a DOCX file and return its content, including text from paragraphs and tables,
    preserving the order of elements.
    """
    doc = docx.Document(docx_path)
    document_paragraphs = doc.paragraphs
    document_tables = doc.tables
    body_elements_order = [("P", paragraph) for paragraph in document_paragraphs] + \
                          [("T", table) for table in document_tables]

    document_text = ""
    for element_type, element in sorted(body_elements_order, key=lambda x: x[1]._element.getparent().index(x[1]._element)):
        if element_type == "P":
            document_text += element.text + "\n"
        elif element_type == "T":
            for row in element.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                document_text += row_text + "\n"

    return document_text.strip()

def case_study_extract(evaluation_info: str) -> tuple:
    """
    Extract specific sections from the evaluation info text using regular expressions.
    Returns a tuple containing instructions, abbreviations, email, and content.
    """
    # Instructions
    instructions = extract_section(evaluation_info, r"(?i)Important Notice[\s\S]+?Please Note", 17, -11)
    # Abbreviations
    abbreviations = extract_section(evaluation_info, r"(?i)Abbreviations[\s\S]+?From:", 0, -5)
    # Email and Content
    email, content = extract_email_and_content(evaluation_info, r"Subject:[\s\S]+?([T/t]hanks?|Head of Unit)")
    
    return instructions, abbreviations, email, content

def extract_section(text: str, pattern: str, start_slice: int, end_slice: int) -> str:
    """Helper function to extract a section from text using a regular expression."""
    match = re.search(pattern, text)
    return match.group(0)[start_slice:end_slice].strip() if match else ""

def extract_email_and_content(text: str, pattern: str) -> tuple:
    """Extract email and remaining content from text using a regular expression."""
    match = re.search(pattern, text)
    if match:
        email = match.group(0).strip()
        content = text[match.end():].strip()
        return email, content
    return "", text