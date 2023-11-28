import pandas as pd
from sqlalchemy.orm import Session
from database.schema import CaseStudyEvaluation
from docx import Document


class Evaluation:
    def __init__(self):
        self.df = pd.read_csv("./training/dataset/test_split.csv")
        self.overall_score = pd.read_csv("./training/results/singleton/summary/overall_score_1031_v3.csv")
        self.communication_score = pd.read_csv(
            "./training/results/singleton/summary/communication_score_1031_v3.csv")
        self.overall_summary = pd.read_csv("./training/results/singleton/summary/overall_summary_1031_v3.csv")
        self.communication_summary = pd.read_csv("./training/results/singleton/summary/communication_summary_1031_v3.csv")
        self.tips_errors = pd.read_csv("./training/results/singleton/summary/tips_errors_1031_v3.csv")

    @staticmethod
    def extract_values(df: pd.DataFrame) -> tuple:
        actual, predicted = df[["Actual", "Predicted"]].values[0]
        return actual, predicted

    def create_evaluation_csv(self, path: str, session: Session) -> None:
        evaluation_csv = []
        for ind, row in self.df.iterrows():
            eval_id = row["Evaluation ID"]
            overall_score_filter = self.extract_values(self.overall_score[self.overall_score["ID"] == eval_id])
            overall_summary_filter = self.extract_values(self.overall_summary[self.overall_summary["ID"] == eval_id])
            communication_score_filter = self.extract_values(self.communication_score[self.communication_score["ID"] == eval_id])
            communication_summary_filter = self.extract_values(self.communication_summary[self.communication_summary["ID"] == eval_id])
            communication_tips_errors_filter = self.extract_values(self.tips_errors[self.tips_errors["ID"] == eval_id])
            trainee_answer_filter = session.get(CaseStudyEvaluation, eval_id).trainee_answer
            evaluation_csv.append([eval_id, overall_score_filter[0], overall_score_filter[1], overall_summary_filter[0],
                                   overall_summary_filter[1],
                                   communication_score_filter[0], communication_score_filter[1],
                                   communication_summary_filter[0], communication_summary_filter[1],
                                   communication_tips_errors_filter[0], communication_tips_errors_filter[1],
                                   trainee_answer_filter])

        evaluation_df = pd.DataFrame(evaluation_csv,
                                     columns=["Evaluation_ID", "Actual Overall Score", "Predicted Overall Score",
                                              "Actual Overall Summary", "Predicted Overall Summary",
                                              "Actual Communication Score", "Predicted Communication Score",
                                              "Actual Communication Summary", "Predicted Communication Summary",
                                              "Actual Tips/Errors", "Predicted Tips/Errors", "Trainee's Answer"])
        evaluation_df.to_csv(path, index=False)

    def create_docx(self, output_path: str, session: Session) -> None:
        for ind, row in self.df.iterrows():
            doc = Document()
            eval_id = row["Evaluation ID"]
            overall_score_filter = self.extract_values(self.overall_score[self.overall_score["ID"] == eval_id])
            overall_summary_filter = self.extract_values(self.overall_summary[self.overall_summary["ID"] == eval_id])
            communication_score_filter = self.extract_values(
                self.communication_score[self.communication_score["ID"] == eval_id])
            communication_summary_filter = self.extract_values(
                self.communication_summary[self.communication_summary["ID"] == eval_id])
            communication_tips_errors_filter = self.extract_values(self.tips_errors[self.tips_errors["ID"] == eval_id])
            trainee_answer_filter = session.get(CaseStudyEvaluation, eval_id).trainee_answer

            doc.add_heading("\nTrainee's Answer\n", 0)
            doc.add_paragraph(trainee_answer_filter)
            doc.add_heading("\nHuman Evaluator\n", 0)
            doc.add_heading("\nActual Overall Score\n", 1)
            doc.add_paragraph(overall_score_filter[0])
            doc.add_heading("\nActual Overall Summary\n", 1)
            doc.add_paragraph(overall_summary_filter[0])
            doc.add_heading("\nActual Communication Score\n", 1)
            doc.add_paragraph(communication_score_filter[0])
            doc.add_heading("\nActual Communication Summary\n", 1)
            doc.add_paragraph(communication_summary_filter[0])
            doc.add_heading("\nActual Tips/Errors\n", 1)
            doc.add_paragraph(communication_tips_errors_filter[0])
            doc.add_heading("\nAI Evaluator\n", 0)
            doc.add_heading("\nPredicted Overall Score\n", 1)
            doc.add_paragraph(overall_score_filter[1])
            doc.add_heading("\nPredicted Overall Summary\n", 1)
            doc.add_paragraph(overall_summary_filter[1])
            doc.add_heading("\nPredicted Communication Score\n", 1)
            doc.add_paragraph(communication_score_filter[1])
            doc.add_heading("\nPredicted Communication Summary\n", 1)
            doc.add_paragraph(communication_summary_filter[1])
            doc.add_heading("\nPredicted Tips/Errors\n", 1)
            doc.add_paragraph(communication_tips_errors_filter[1])

            doc.save(f'{output_path}/evaluation_{eval_id}.docx')


evaluation_csv_docx = Evaluation()
