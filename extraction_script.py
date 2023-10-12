import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import re


def read_docx(docx_path):
    doc = docx.Document(docx_path)
    document_text = ""
    document_paragraph = doc.paragraphs
    document_table = doc.tables
    body_location = []
    for element in doc.element.body:
        if isinstance(element, CT_Tbl):
            body_location.append("T")
        elif isinstance(element, CT_P):
            body_location.append("P")
    para, tab = 0, 0
    for loc in body_location:
        if loc == "P":
            document_text += document_paragraph[para].text + "\n"
            para += 1
        elif loc == "T":
            for row in document_table[tab].rows:
                row_text = ""
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text
                        row_text += paragraph_text + " | "
                document_text += row_text.strip() + "\n"
            tab += 1
    document_text = document_text.strip()
    return document_text


def trainee_answer_extractor(trainee_answer):
    question_start = trainee_answer.find("Question")
    trainee_answer_start = trainee_answer.find("Trainee's Answer")
    case_study_name = trainee_answer[question_start:trainee_answer_start].strip()
    answer_content = trainee_answer[trainee_answer_start:].strip()
    return case_study_name, answer_content


def review_guide_extract(review_info):
    pattern = r"Communication[\s\S]+?Key tips to improve / maintain performance:"
    communication_section = re.search(pattern, review_info)
    recommendations_solution_pattern = r"Example [S/s]olution"
    recommendations_solution_search = re.search(recommendations_solution_pattern, review_info).end()
    recommendations_solution = review_info[recommendations_solution_search:]
    score_grid = ""
    if communication_section is not None:
        score_grid = communication_section.group(0)
    return score_grid, recommendations_solution


def evaluation_info_extract(evaluation_info):
    # Overall Score
    score_pattern = "Your [S/s]core .*"
    overall_score = re.search(score_pattern, evaluation_info).group(0).strip()
    overall_score = re.findall(r"\d+\.\d+|\d+", overall_score)
    overall_score = float(overall_score[0])
    # Overall Summary
    summary_pattern = r"Summary[\s\S]+?(Communication|Per Competency Score?)"
    summary = re.search(summary_pattern, evaluation_info).group(0).strip()
    # Communication Score
    communication_pattern = r"Communication .*"
    communication_search = re.search(communication_pattern, evaluation_info)
    communication_score = communication_search.group(0).strip()
    communication_score = re.findall(r"\d+\.\d+|\d+", communication_score)
    communication_score = float(communication_score[0])
    # Tips to Improve
    tips_pattern = r"(?i)(Tips to Improve)"
    tips_search = re.search(tips_pattern, evaluation_info)
    errors_search = None
    communication_summary = ""
    errors_info = ""
    tips_to_improve = ""
    if tips_search is not None:
        tips_to_improve = evaluation_info[tips_search.start():].strip()
    else:
        # Errors Info
        errors_pattern = r"(?:Spelling|Grammar)"
        errors_search = re.search(errors_pattern, evaluation_info)
    # Communication Summary
    if communication_search is not None:
        start = communication_search.end()
        if tips_search is not None:
            communication_summary = evaluation_info[start: tips_search.start()].strip()
        elif errors_search is not None:
            communication_summary = evaluation_info[start: errors_search.start()].strip()
            errors_info = evaluation_info[errors_search.start():].strip()
    return overall_score, summary, communication_score, communication_summary, errors_info, tips_to_improve


def case_study_extract(evaluation_info):
    # Instructions
    important_notice = r"(?i)Important Notice[\s\S]+?Please Note"
    instructions_search = re.search(important_notice, evaluation_info)
    instructions = ""
    if instructions_search is not None:
        instructions = instructions_search.group(0)[17:-11].strip()
    # Abbreviations
    abb_pattern = r"(?i)Abbreviations[\s\S]+?From:"
    abb_search = re.search(abb_pattern, evaluation_info)
    abbreviations = ""
    if abb_search is not None:
        abbreviations = abb_search.group(0)[:-5].strip()
    # Email and Content
    email_pattern = r"Subject:[\s\S]+?([T/t]hanks?|Head of Unit)"
    email_search = re.search(email_pattern, evaluation_info)
    email = ""
    content = evaluation_info
    if email_search is not None:
        email = email_search.group(0).strip()
        content = evaluation_info[email_search.end():].strip()
    return instructions, abbreviations, email, content


# print(evaluation_info_extract(read_docx(r"E:\Freelancing\AXEOM\Axeom_EUTraining\CS docs for AI\Generic Case Studies\Case 1\Copy of 2_2_Evaluation_DT.docx")))
